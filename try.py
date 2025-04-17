import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

# 이미지 폴더 경로
img_folder = r"C:\Users\toozi\OneDrive\문서\GitHub\BizLab\img"

# 해당 폴더 안의 .png 파일 목록을 가져와 img_path 리스트로 설정
img_path = [os.path.join(img_folder, fname) 
            for fname in sorted(os.listdir(img_folder)) 
            if fname.endswith(".png")]

json_path = r"C:\Users\toozi\OneDrive\문서\GitHub\BizLab\data\survey_result.json"
with open(json_path, 'r', encoding='utf-8-sig') as f:
    survey_json = json.load(f)

print(img_path)

'''doc = Document()
# 제목
doc.add_heading("ERG 이론에 따른 대학생의 SNS 사용 동기와 SNS 사용 만족도간의 관계 연구 설문조사 보고서", level=0)

# 목차
doc.add_page_break()
doc.add_heading("목차", level=1)
sections = ["[1] 설문 주제 및 조사 개요", "[2] Introduction", "[3] 단일 질문 분석", "[4] 교차 분석", "[5] 결론", "[6] 기타"]
for section in sections:
    doc.add_paragraph(section)

# [1] 설문 주제 및 조사 개요
doc.add_page_break()
doc.add_heading("[1] 설문 주제 및 조사 개요", level=1)

# [2] Introduction
doc.add_page_break()
doc.add_heading("[2] Introduction", level=1)

# [3] 단일 질문 분석
doc.add_page_break()
doc.add_heading("[3] 단일 질문 분석", level=1)

# [4] 교차 분석
doc.add_page_break()
doc.add_heading("[4] 교차 분석", level=1)

# [5] 결론
doc.add_page_break()
doc.add_heading("[5] 결론", level=1)

# [6] 기타
doc.add_page_break()
doc.add_heading("[6] 기타", level=1)

# 저장
output_name = "example.docx"
doc.save(output_name)
print(f"✅ Word 저장 완료: {output_name}")'''