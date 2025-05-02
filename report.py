import re 
import os
from docx import Document
from docx.shared import Inches, Pt 

def survey_report(title, sections, outline_data, img_path, survey_content, cross_img_folder, cross_analy_text, local, output_name):

    doc = Document()

    # 1) 제목
    doc.add_heading(title, level=0)

    # 2) 목차
    doc.add_page_break()
    doc.add_heading("목차", level=1)
    doc.add_paragraph("\n")

    for section in sections:
        para = doc.add_paragraph()
        run = para.add_run(section)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(11)

    # 3) 내용
    ## [1] 설문 주제 및 조사 개요
    doc.add_page_break()
    doc.add_heading("[1] 설문 주제 및 조사 개요", level=1)
    doc.add_heading("1. 조사목적", level=2)
    
    para1 = doc.add_paragraph()
    run1 = para1.add_run(f"• {outline_data['조사목적']}")
    run1.font.name = '맑은 고딕'
    run1.font.size = Pt(11)

    doc.add_heading("2. 조사설계 및 방법", level=2)
    outline_contents = ["조사대상", "조사기간", "조사방법", "조사내용", "참여인원"]

    for content in outline_contents:
        value = outline_data.get(content, "정보 없음")
        para = doc.add_paragraph()
        run = para.add_run(f"• {content} : {value}")
        run.font.name = '맑은 고딕'
        run.font.size = Pt(11)

    ## [2] 단일 질문 분석
    doc.add_page_break()
    doc.add_heading("[2] 단일 질문 분석", level=1)

    for i, img in enumerate(img_path):
        doc.add_heading(f"질문 {i+1}", level=2)

        # 이미지 삽입
        try:
            doc.add_picture(img, width=Inches(5))
        except Exception as e:
            doc.add_paragraph(f"이미지 삽입 오류: {e}")

        # 해석 텍스트 삽입
        try:
            analysis_text = survey_content[i] if i < len(survey_content) else "해석 없음"
        except Exception as e:
            analysis_text = f"해석 실패: {e}"

        # 문장 단위로 줄바꿈 처리
        sentences = re.split(r'(?<=[.!?])\s+', analysis_text.strip())  # 마침표, 느낌표, 물음표 뒤 기준
        for sentence in sentences:
            para = doc.add_paragraph()
            run = para.add_run(sentence.strip())
            run.font.name = '맑은 고딕'
            run.font.size = Pt(11)

        doc.add_paragraph("\n")  # 간격 띄우기

    ## [3] 교차 분석
    doc.add_page_break()
    doc.add_heading("[3] 교차 분석", level=1)

    cross_img_paths = [os.path.join(cross_img_folder, fname) 
                 for fname in sorted(os.listdir(cross_img_folder)) 
                 if fname.lower().endswith(('.png', '.jpg', '.jpeg'))]


    for i, cross_path in enumerate(cross_img_paths):
        doc.add_heading(f"교차 분석 결과 {i+1}", level=2)

        # 이미지 삽입
        try:
            doc.add_picture(cross_path, width=Inches(5))
        except Exception as e:
            doc.add_paragraph(f"이미지 삽입 오류: {e}")

        # 해석 텍스트 삽입
        try:
            # cross_analy_text[i]를 사용하여 분석 텍스트 가져오기
            analysis_text = cross_analy_text[i] if i < len(cross_analy_text) else "해석 없음"
        except Exception as e:
            analysis_text = f"해석 실패: {e}"

        # 문장 단위 줄바꿈
        sentences = re.split(r'(?<=[.!?])\s+', analysis_text.strip())
        for sentence in sentences:
            para = doc.add_paragraph()
            run = para.add_run(sentence.strip())
            run.font.name = '맑은 고딕'
            run.font.size = Pt(11)

        doc.add_paragraph("\n")  # 간격 띄우기


    # 4) 저장
    save_doc = os.path.join(local, f'{output_name}')
    
    doc.save(save_doc)
    print(f" Word 저장 완료: {output_name}")



## gpt한테 index 받아오는 경우 
'''
def add_index_entry(text, indent_level=0): 
    para = doc.add_paragraph()
    run = para.add_run("    " * indent_level + text)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(11)

# parsed_index에 따라 문서에 항목 추가
for entry in index_content:
    title, subentries = entry
    add_index_entry(title, indent_level=0)
    for subentry in subentries:
        add_index_entry(subentry, indent_level=1)
'''